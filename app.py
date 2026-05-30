from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
import os
import io
import json
import html
from groq import Groq
from supabase import create_client, Client
from functools import wraps
from dotenv import load_dotenv
from flask_cors import CORS
import ast
import subprocess
import uuid
import shutil
import re
import base64
import hashlib
import hmac
from datetime import datetime, timezone
import urllib.error
import urllib.parse
import urllib.request
from html.parser import HTMLParser
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

env_path = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(env_path)

def clear_dead_local_proxy():
    proxy_vars = (
        "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY",
        "http_proxy", "https_proxy", "all_proxy",
        "GIT_HTTP_PROXY", "GIT_HTTPS_PROXY",
    )
    for name in proxy_vars:
        value = os.environ.get(name, "")
        if "127.0.0.1:9" in value or "localhost:9" in value:
            os.environ.pop(name, None)

clear_dead_local_proxy()

# --- FIX 1: Explicit Static Folder for Generated Outputs ---
app = Flask(__name__, static_url_path='', static_folder='static')
# Standardizing CORS to allow the frontend to talk to these specific routes
CORS(app, resources={r"/api/*": {"origins": "*"}})

OUTPUTS_DIR = os.path.join(app.root_path, 'static', 'outputs')
os.makedirs('static/css', exist_ok=True)
os.makedirs('static/js', exist_ok=True)
os.makedirs('templates', exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)

api_key = os.environ.get("GROQ_API_KEY", "").strip()
client = Groq(api_key=api_key) if api_key else None
MODEL_ID = 'llama-3.3-70b-versatile' 
FREE_CREDIT_ALLOWANCE = 400
EXPERIMENT_CREDIT_COST = 20

SUPABASE_URL = os.environ.get("SUPABASE_URL", "").strip()
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "").strip()
SUPABASE_SERVICE_ROLE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "").strip()
supabase = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL else None

RAZORPAY_KEY_ID = os.environ.get("RAZORPAY_KEY_ID", "").strip()
RAZORPAY_KEY_SECRET = os.environ.get("RAZORPAY_KEY_SECRET", "").strip()
RAZORPAY_WEBHOOK_SECRET = os.environ.get("RAZORPAY_WEBHOOK_SECRET", "").strip()
RAZORPAY_CREDIT_PACK_AMOUNT_INR = int(os.environ.get("RAZORPAY_CREDIT_PACK_AMOUNT_INR", "99") or 99)
RAZORPAY_CREDIT_PACK_CREDITS = int(os.environ.get("RAZORPAY_CREDIT_PACK_CREDITS", "400") or 400)

PAYMENT_LINK_URL = (
    os.environ.get("RAZORPAY_PAYMENT_LINK_URL", "").strip()
    or os.environ.get("STRIPE_PAYMENT_LINK_URL", "").strip()
    or os.environ.get("LIBRA_PAYMENT_LINK_URL", "").strip()
)

def parse_model_json(raw_text, default_key="content"):
    if not raw_text:
        return {default_key: ""}

    raw_text = raw_text.strip()

    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        pass

    fenced_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", raw_text, re.DOTALL)
    if fenced_match:
        try:
            return json.loads(fenced_match.group(1))
        except json.JSONDecodeError:
            pass

    object_match = re.search(r"\{.*\}", raw_text, re.DOTALL)
    if object_match:
        try:
            return json.loads(object_match.group(0))
        except json.JSONDecodeError:
            pass

    return {default_key: raw_text}

def is_code_section(section_title):
    title = (section_title or "").strip().lower()
    return "code" in title or "program" in title or "implementation" in title

def is_output_section(section_title):
    title = (section_title or "").strip().lower()
    return "result" in title or "output" in title or "observation" in title

def strip_wrapping_fences(text):
    if not text:
        return ""

    value = text.strip()
    match = re.fullmatch(r"```(?:json)?\s*([\s\S]*?)\s*```", value)
    return match.group(1).strip() if match else value

def clean_generated_section_content(content, section_title, preserve_code=False):
    if content is None:
        return ""

    text = str(content).strip()
    if not text:
        return ""

    if not preserve_code:
        text = strip_wrapping_fences(text)

    section = re.escape((section_title or "").strip())
    if section:
        patterns = [
            rf"^\s*#{1,6}\s*{section}\s*:?\s*",
            rf"^\s*\*\*{section}\s*:?\*\*\s*",
            rf"^\s*<h[1-6][^>]*>\s*{section}\s*:?\s*</h[1-6]>\s*",
            rf"^\s*<strong[^>]*>\s*{section}\s*:?\s*</strong>\s*",
            rf"^\s*{section}\s*:\s*",
        ]
        for pattern in patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    return text.strip()

def get_section_specific_guidance(section_title, subject_type, is_visual):
    title = (section_title or "").strip().lower()

    if is_visual:
        if subject_type == "Coding":
            return (
                "This is a coding visual/code section. Return only the code needed for this section. "
                "Do not include introductions, explanations, labels, JSON wrappers, or markdown outside the code block."
            )
        return (
            "This is a non-coding visual section. Return only the requested diagram/visual content for this section. "
            "Do not include introductions, explanations, or unrelated practical sections."
        )

    if "aim" in title or "objective" in title:
        return (
            "SECTION TYPE: AIM. Return exactly one concise aim/objective statement for this experiment. "
            "Do not include steps, theory, code, output, datasets, implementation details, or extra headings."
        )
    if "theory" in title or "concept" in title:
        return (
            "SECTION TYPE: THEORY. Explain only the core concepts, definitions, formulas, and academic background needed for this experiment. "
            "Do not include aim, procedure, implementation steps, source code, output, result, or conclusion."
        )
    if "algorithm" in title or "procedure" in title or "method" in title:
        return (
            "SECTION TYPE: PROCEDURE/ALGORITHM. Return only the stepwise method or algorithm. "
            "Use a compact numbered list. Do not include theory paragraphs, code, output, result, or conclusion."
        )
    if "code" in title or "program" in title or "implementation" in title:
        return (
            "SECTION TYPE: CODE. Return only the final runnable source code for the experiment in one fenced code block. "
            "Do not include aim, theory, explanation, algorithm prose, output, result, JSON wrappers, or repeated section labels. "
            "Ensure the generated JSON is strictly valid. Do not use triple quotes inside the JSON object."
        )
    if "result" in title or "output" in title or "observation" in title:
        if subject_type == "Coding":
            return (
                "SECTION TYPE: CODE OUTPUT. Return strictly only the exact expected console output or generated-result text for the program. "
                "Do not explain the output. Do not include aim, theory, algorithm, source code, conclusion, labels, or markdown headings."
            )
        return (
            "SECTION TYPE: OUTPUT/RESULT. Return only the expected output, observations, or result statement for this experiment. "
            "Do not include aim, theory, algorithm, source code, conclusion, labels, or markdown headings."
        )
    if "conclusion" in title:
        return (
            "SECTION TYPE: CONCLUSION. Return only a short conclusion based on the experiment. "
            "Do not include theory, procedure, algorithm, code, output tables, or new explanatory sections."
        )

    return (
        f"This section must contain only content appropriate for '{section_title}'. "
        "Do not drift into other practical sections or repeat a generic introduction."
    )

# Backend Auth JWT Verification Decorator
def require_auth(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not supabase:
            return jsonify({"error": "Authentication is not configured. Add SUPABASE_URL and SUPABASE_KEY."}), 503
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            return jsonify({"error": "Unauthorized"}), 401
        token = auth_header.split(" ")[1]
        try:
            user = supabase.auth.get_user(token)
            if not user:
                raise Exception("Invalid Session")
            request.user = user
            request.access_token = token
        except Exception as e:
            return jsonify({"error": "Authentication Failed. " + str(e)}), 401
        return f(*args, **kwargs)
    return decorated_function

def get_auth_user_id():
    return request.user.user.id

def get_auth_access_token():
    return getattr(request, "access_token", None)

def supabase_rest_request(method, path, token, body=None):
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise RuntimeError("Supabase is not configured.")
    if not token:
        raise RuntimeError("Missing user session token.")

    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = None
    headers = {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {token}",
        "Accept": "application/json",
    }

    if body is not None:
        payload = json.dumps(body).encode("utf-8")
        headers["Content-Type"] = "application/json"
        headers["Prefer"] = "return=representation"

    req = urllib.request.Request(url, data=payload, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=20) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else []
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            details = json.loads(raw)
        except json.JSONDecodeError:
            details = {"message": raw or exc.reason, "code": exc.code}
        raise RuntimeError(details) from exc

def supabase_service_request(method, path, body=None):
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise RuntimeError("Supabase service role key is required for billing automation.")

    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = None
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Accept": "application/json",
    }

    if body is not None:
        payload = json.dumps(body).encode("utf-8")
        headers["Content-Type"] = "application/json"
        headers["Prefer"] = "return=representation"

    req = urllib.request.Request(url, data=payload, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=20) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else []
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            details = json.loads(raw)
        except json.JSONDecodeError:
            details = {"message": raw or exc.reason, "code": exc.code}
        raise RuntimeError(details) from exc

def razorpay_orders_enabled():
    return bool(RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET and SUPABASE_SERVICE_ROLE_KEY)

def razorpay_api_request(method, path, body=None):
    if not RAZORPAY_KEY_ID or not RAZORPAY_KEY_SECRET:
        raise RuntimeError("Razorpay Orders are not configured.")

    credentials = f"{RAZORPAY_KEY_ID}:{RAZORPAY_KEY_SECRET}".encode("utf-8")
    headers = {
        "Authorization": "Basic " + base64.b64encode(credentials).decode("ascii"),
        "Accept": "application/json",
    }
    payload = None
    if body is not None:
        payload = json.dumps(body).encode("utf-8")
        headers["Content-Type"] = "application/json"

    req = urllib.request.Request(f"https://api.razorpay.com/v1/{path}", data=payload, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=20) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            details = json.loads(raw)
        except json.JSONDecodeError:
            details = {"message": raw or exc.reason, "code": exc.code}
        raise RuntimeError(details) from exc

def verify_hmac_sha256(message, signature, secret):
    expected = hmac.new(secret.encode("utf-8"), message, hashlib.sha256).hexdigest()
    return hmac.compare_digest(expected, signature or "")

def get_billing_order(razorpay_order_id):
    order_filter = urllib.parse.quote(str(razorpay_order_id), safe="")
    rows = supabase_service_request(
        "GET",
        f"billing_orders?select=*&razorpay_order_id=eq.{order_filter}&limit=1"
    )
    return rows[0] if rows else None

def add_credits_with_service(user_id, credits_to_add):
    user_filter = urllib.parse.quote(str(user_id), safe="")
    rows = supabase_service_request(
        "GET",
        f"user_credits?select=credits&user_id=eq.{user_filter}&limit=1"
    )
    if rows:
        current = int(rows[0].get("credits", 0) or 0)
        supabase_service_request(
            "PATCH",
            f"user_credits?user_id=eq.{user_filter}",
            {"credits": current + int(credits_to_add)}
        )
    else:
        supabase_service_request(
            "POST",
            "user_credits",
            {"user_id": user_id, "credits": FREE_CREDIT_ALLOWANCE + int(credits_to_add)}
        )

def fulfill_billing_order(razorpay_order_id, razorpay_payment_id=None, webhook_event_id=None):
    order = get_billing_order(razorpay_order_id)
    if not order:
        raise RuntimeError(f"Unknown Razorpay order: {razorpay_order_id}")
    if order.get("status") != "created":
        return False, order

    order_filter = urllib.parse.quote(str(razorpay_order_id), safe="")
    claimed = supabase_service_request(
        "PATCH",
        f"billing_orders?razorpay_order_id=eq.{order_filter}&status=eq.created",
        {
            "status": "processing",
            "razorpay_payment_id": razorpay_payment_id,
            "webhook_event_id": webhook_event_id
        }
    )
    if not claimed:
        latest = get_billing_order(razorpay_order_id)
        return False, latest or order

    add_credits_with_service(order["user_id"], int(order.get("credits", 0) or 0))
    updated = supabase_service_request(
        "PATCH",
        f"billing_orders?razorpay_order_id=eq.{order_filter}",
        {
            "status": "paid",
            "razorpay_payment_id": razorpay_payment_id,
            "webhook_event_id": webhook_event_id,
            "paid_at": datetime.now(timezone.utc).isoformat()
        }
    )
    return True, updated[0] if updated else order

def credit_setup_error(error):
    text = str(error).lower()
    if "row-level security" in text or "42501" in text:
        return (
            "Credits table is missing the insert policy. Run the updated "
            "user_credits SQL setup in Supabase, then reload Labmate.ai."
        )
    return None

def ensure_credit_account(user_id):
    token = get_auth_access_token()
    user_filter = urllib.parse.quote(str(user_id), safe="")
    rows = supabase_rest_request(
        "GET",
        f"user_credits?select=credits&user_id=eq.{user_filter}",
        token
    )
    if rows:
        credits = int(rows[0].get("credits", 0) or 0)
        return credits, False

    try:
        supabase_rest_request(
            "POST",
            "user_credits",
            token,
            {"user_id": user_id, "credits": FREE_CREDIT_ALLOWANCE}
        )
    except RuntimeError as e:
        if "23505" in str(e):
            rows = supabase_rest_request(
                "GET",
                f"user_credits?select=credits&user_id=eq.{user_filter}",
                token
            )
            if rows:
                return int(rows[0].get("credits", 0) or 0), False
        raise
    return FREE_CREDIT_ALLOWANCE, True

def get_current_credits(user_id):
    credits, _created = ensure_credit_account(user_id)
    return credits

class SimpleHTMLNode:
    def __init__(self, tag="root", attrs=None):
        self.tag = tag
        self.attrs = dict(attrs or [])
        self.children = []

class SimpleHTMLTreeParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = SimpleHTMLNode()
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        node = SimpleHTMLNode(tag.lower(), attrs)
        self.stack[-1].children.append(node)
        if tag.lower() not in {"br", "img", "hr", "meta", "link", "input"}:
            self.stack.append(node)

    def handle_endtag(self, tag):
        tag = tag.lower()
        while len(self.stack) > 1:
            node = self.stack.pop()
            if node.tag == tag:
                break

    def handle_data(self, data):
        if data:
            self.stack[-1].children.append(data)

def parse_html_fragment(fragment):
    parser = SimpleHTMLTreeParser()
    parser.feed(fragment or "")
    return parser.root

def node_text(node):
    if isinstance(node, str):
        return node
    if node.tag == "br":
        return "\n"
    return "".join(node_text(child) for child in node.children)

def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.footer_distance = Inches(0.35)
    section.header_distance = Inches(0.35)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "64748B")

def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def style_run(run, bold=False, code=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold

def add_inline_runs(paragraph, node, bold=False, code=False):
    if isinstance(node, str):
        text = re.sub(r"\s+", " ", node)
        if text:
            run = paragraph.add_run(text)
            style_run(run, bold=bold, code=code)
        return
    if node.tag == "br":
        paragraph.add_run().add_break()
        return
    child_bold = bold or node.tag in {"strong", "b"}
    child_code = code or node.tag in {"code", "pre"}
    for child in node.children:
        add_inline_runs(paragraph, child, child_bold, child_code)

def add_text_paragraph(container, text="", bold=False, code=False):
    if not text.strip():
        return None
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text.strip())
    style_run(run, bold=bold, code=code)
    return paragraph

def render_table(container, table_node):
    rows = [child for child in table_node.children if not isinstance(child, str) and child.tag in {"tr", "thead", "tbody"}]
    flat_rows = []
    for row in rows:
        if row.tag == "tr":
            flat_rows.append(row)
        else:
            flat_rows.extend(child for child in row.children if not isinstance(child, str) and child.tag == "tr")
    if not flat_rows:
        return
    max_cols = max(
        len([cell for cell in row.children if not isinstance(cell, str) and cell.tag in {"td", "th"}])
        for row in flat_rows
    )
    if max_cols < 1:
        return
    try:
        table = container.add_table(rows=len(flat_rows), cols=max_cols)
    except TypeError:
        table = container.add_table(rows=len(flat_rows), cols=max_cols, width=Inches(6.5))
    table.autofit = True
    for row_index, row in enumerate(flat_rows):
        cells = [cell for cell in row.children if not isinstance(cell, str) and cell.tag in {"td", "th"}]
        for col_index, cell_node in enumerate(cells):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, cell_node, bold=(cell_node.tag == "th"))

def render_html_fragment(container, fragment):
    root = parse_html_fragment(fragment)
    for child in root.children:
        render_node(container, child)

def render_node(container, node):
    if isinstance(node, str):
        add_text_paragraph(container, node)
        return
    if node.tag in {"script", "style", "button", "svg"}:
        return
    if node.tag == "table":
        render_table(container, node)
        return
    if node.tag == "pre":
        add_text_paragraph(container, node_text(node), code=True)
        return
    if node.tag in {"p", "div", "section", "article", "li", "h1", "h2", "h3", "h4", "h5", "h6"}:
        text = node_text(node).strip()
        has_block_child = any(
            not isinstance(child, str) and child.tag in {"table", "p", "div", "pre", "ul", "ol"}
            for child in node.children
        )
        if text and not has_block_child:
            paragraph = container.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1
            add_inline_runs(paragraph, node, bold=node.tag.startswith("h"))
        else:
            for child in node.children:
                render_node(container, child)
        return
    if node.tag in {"ul", "ol"}:
        for child in node.children:
            if not isinstance(child, str) and child.tag == "li":
                add_text_paragraph(container, node_text(child), bold=False)
        return
    text = node_text(node).strip()
    if text:
        add_text_paragraph(container, text)

def build_docx_export(title, document_data):
    document = Document()
    set_document_defaults(document)

    section = document.sections[0]
    if document_data.get("header_html"):
        render_html_fragment(section.header, document_data.get("header_html"))

    footer_text = node_text(parse_html_fragment(document_data.get("footer_html", ""))).strip()
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
    if footer_text:
        run = paragraph.add_run(footer_text)
        style_run(run)
    paragraph.add_run("\t")
    add_page_field(paragraph)

    for item in document_data.get("sections", []):
        title_text = (item.get("title") or "").strip()
        if title_text:
            add_text_paragraph(document, title_text, bold=True)
        render_html_fragment(document, item.get("body_html", ""))

    return document

@app.route('/')
def index():
    # Inject secure config to the frontend JS engine
    return render_template('index.html', supabase_url=SUPABASE_URL, supabase_key=SUPABASE_KEY)

@app.route('/favicon.ico')
def favicon():
    return '', 204

# --- FIX 2: Added a helper to serve generated images without 404s ---
@app.route('/outputs/<path:path>')
def send_report_assets(path):
    return send_from_directory(OUTPUTS_DIR, path)

@app.route('/api/outline', methods=['POST'])
@require_auth
def generate_outline():
    # ... (Keep your current logic, it is structurally sound)
    try:
        if not client:
            return jsonify({"error": "Groq API key is missing from the .env file. Please add GROQ_API_KEY."}), 500
            
        data = request.json
        prompt_text = data.get('prompt', '')
        response = client.chat.completions.create(
            model=MODEL_ID,
            messages=[
                {"role": "system", "content": "You are Labmate.ai. Generate a professional lab practical outline. You must output a strictly valid JSON object with a single key 'outline' containing an array of strings representing the section titles."},
                {"role": "user", "content": f"Generate a professional lab practical outline for: {prompt_text}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.2
        )
        content = json.loads(response.choices[0].message.content)
        return jsonify(content.get('outline', []))
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/credits/account', methods=['GET'])
@require_auth
def credit_account():
    try:
        user_id = get_auth_user_id()
        credits, created = ensure_credit_account(user_id)
        return jsonify({
            "credits": credits,
            "created": created,
            "free_allowance": FREE_CREDIT_ALLOWANCE,
            "experiment_cost": EXPERIMENT_CREDIT_COST,
            "free_experiments": FREE_CREDIT_ALLOWANCE // EXPERIMENT_CREDIT_COST
        })
    except Exception as e:
        setup_message = credit_setup_error(e)
        if setup_message:
            return jsonify({"error": "CREDITS_SETUP_REQUIRED", "message": setup_message}), 500
        return jsonify({"error": str(e)}), 500

@app.route('/api/credits/consume_experiment', methods=['POST'])
@require_auth
def consume_experiment_credits():
    try:
        user_id = get_auth_user_id()
        current_credits = get_current_credits(user_id)
        if current_credits < EXPERIMENT_CREDIT_COST:
            return jsonify({
                "error": "OUT_OF_CREDITS",
                "message": "You need more credits to generate another experiment.",
                "credits": current_credits,
                "required": EXPERIMENT_CREDIT_COST
            }), 402

        remaining = current_credits - EXPERIMENT_CREDIT_COST
        user_filter = urllib.parse.quote(str(user_id), safe="")
        supabase_rest_request(
            "PATCH",
            f"user_credits?user_id=eq.{user_filter}",
            get_auth_access_token(),
            {"credits": remaining}
        )
        return jsonify({
            "credits": remaining,
            "charged": EXPERIMENT_CREDIT_COST
        })
    except Exception as e:
        setup_message = credit_setup_error(e)
        if setup_message:
            return jsonify({"error": "CREDITS_SETUP_REQUIRED", "message": setup_message}), 500
        return jsonify({"error": str(e)}), 500

@app.route('/api/billing/options', methods=['GET'])
@require_auth
def billing_options():
    orders_ready = razorpay_orders_enabled()
    return jsonify({
        "methods": ["Cards", "UPI", "Net banking", "Wallets"],
        "configured": orders_ready or bool(PAYMENT_LINK_URL),
        "mode": "orders" if orders_ready else "payment_link",
        "key_id": RAZORPAY_KEY_ID if orders_ready else None,
        "amount_inr": RAZORPAY_CREDIT_PACK_AMOUNT_INR,
        "credits": RAZORPAY_CREDIT_PACK_CREDITS,
        "checkout_url": PAYMENT_LINK_URL if PAYMENT_LINK_URL else None,
        "message": (
            f"Buy {RAZORPAY_CREDIT_PACK_CREDITS} credits for INR {RAZORPAY_CREDIT_PACK_AMOUNT_INR}."
            if orders_ready
            else (
                "Billing is ready to open your configured payment link."
                if PAYMENT_LINK_URL
                else "Add Razorpay Orders keys or RAZORPAY_PAYMENT_LINK_URL to enable live payments."
            )
        )
    })

@app.route('/api/billing/create_order', methods=['POST'])
@require_auth
def billing_create_order():
    try:
        if not razorpay_orders_enabled():
            return jsonify({
                "error": "RAZORPAY_ORDERS_NOT_CONFIGURED",
                "message": "Add RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET, and SUPABASE_SERVICE_ROLE_KEY to enable Razorpay Orders."
            }), 501

        user_id = get_auth_user_id()
        amount_paise = RAZORPAY_CREDIT_PACK_AMOUNT_INR * 100
        receipt = f"lm_{uuid.uuid4().hex[:24]}"
        order = razorpay_api_request("POST", "orders", {
            "amount": amount_paise,
            "currency": "INR",
            "receipt": receipt,
            "notes": {
                "user_id": str(user_id),
                "credits": str(RAZORPAY_CREDIT_PACK_CREDITS),
                "product": "labmate_credits"
            }
        })

        supabase_service_request("POST", "billing_orders", {
            "user_id": user_id,
            "razorpay_order_id": order["id"],
            "amount_paise": amount_paise,
            "currency": "INR",
            "credits": RAZORPAY_CREDIT_PACK_CREDITS,
            "status": "created",
            "receipt": receipt
        })

        return jsonify({
            "key_id": RAZORPAY_KEY_ID,
            "order_id": order["id"],
            "amount": amount_paise,
            "currency": "INR",
            "credits": RAZORPAY_CREDIT_PACK_CREDITS,
            "name": "Labmate.ai",
            "description": f"{RAZORPAY_CREDIT_PACK_CREDITS} Labmate.ai credits"
        })
    except Exception as e:
        return jsonify({"error": "ORDER_CREATE_FAILED", "message": str(e)}), 500

@app.route('/api/billing/verify', methods=['POST'])
@require_auth
def billing_verify_payment():
    try:
        if not RAZORPAY_KEY_SECRET:
            return jsonify({"error": "RAZORPAY_NOT_CONFIGURED", "message": "Razorpay key secret is missing."}), 501

        data = request.json or {}
        razorpay_order_id = data.get("razorpay_order_id", "")
        razorpay_payment_id = data.get("razorpay_payment_id", "")
        razorpay_signature = data.get("razorpay_signature", "")
        if not all([razorpay_order_id, razorpay_payment_id, razorpay_signature]):
            return jsonify({"error": "INVALID_PAYMENT_RESPONSE", "message": "Missing Razorpay payment details."}), 400

        order = get_billing_order(razorpay_order_id)
        if not order or order.get("user_id") != get_auth_user_id():
            return jsonify({"error": "UNKNOWN_ORDER", "message": "This payment order was not found for your account."}), 404

        message = f"{razorpay_order_id}|{razorpay_payment_id}".encode("utf-8")
        if not verify_hmac_sha256(message, razorpay_signature, RAZORPAY_KEY_SECRET):
            return jsonify({"error": "SIGNATURE_VERIFICATION_FAILED", "message": "Payment signature verification failed."}), 400

        credited, updated_order = fulfill_billing_order(razorpay_order_id, razorpay_payment_id=razorpay_payment_id)
        credits, _created = ensure_credit_account(get_auth_user_id())
        return jsonify({
            "status": "paid",
            "credited": credited,
            "credits_added": int(updated_order.get("credits", 0) or 0),
            "credits": credits
        })
    except Exception as e:
        return jsonify({"error": "PAYMENT_VERIFY_FAILED", "message": str(e)}), 500

@app.route('/api/razorpay/webhook', methods=['POST'])
def razorpay_webhook():
    try:
        if not RAZORPAY_WEBHOOK_SECRET:
            return jsonify({"error": "Webhook secret is not configured."}), 501

        raw_body = request.get_data() or b""
        signature = request.headers.get("X-Razorpay-Signature", "")
        if not verify_hmac_sha256(raw_body, signature, RAZORPAY_WEBHOOK_SECRET):
            return jsonify({"error": "Invalid webhook signature."}), 400

        payload = json.loads(raw_body.decode("utf-8"))
        event = payload.get("event", "")
        if event not in {"order.paid", "payment.captured"}:
            return jsonify({"status": "ignored", "event": event})

        payment_entity = (payload.get("payload", {}).get("payment", {}) or {}).get("entity", {}) or {}
        order_entity = (payload.get("payload", {}).get("order", {}) or {}).get("entity", {}) or {}
        razorpay_order_id = payment_entity.get("order_id") or order_entity.get("id")
        razorpay_payment_id = payment_entity.get("id")
        if not razorpay_order_id:
            return jsonify({"error": "Webhook payload did not include an order id."}), 400

        credited, _order = fulfill_billing_order(
            razorpay_order_id,
            razorpay_payment_id=razorpay_payment_id,
            webhook_event_id=request.headers.get("X-Razorpay-Event-Id")
        )
        return jsonify({"status": "ok", "credited": credited})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/billing/checkout', methods=['POST'])
@require_auth
def billing_checkout():
    if not PAYMENT_LINK_URL:
        return jsonify({
            "error": "BILLING_NOT_CONFIGURED",
            "message": "Create a Razorpay/Stripe payment link and add it to .env as RAZORPAY_PAYMENT_LINK_URL or STRIPE_PAYMENT_LINK_URL."
        }), 501
    return jsonify({"checkout_url": PAYMENT_LINK_URL})

@app.route('/api/export/word', methods=['POST'])
@require_auth
def export_word():
    try:
        data = request.json or {}
        title = re.sub(r"[^A-Za-z0-9_-]+", "_", data.get("title", "Labmate.ai_experiment")).strip("_") or "Labmate.ai_experiment"
        document_data = data.get("document") or {}
        document_html = data.get("html", "")
        has_structured_content = bool(document_data.get("sections") or document_data.get("header_html") or document_data.get("footer_html"))
        if not has_structured_content and not document_html.strip():
            return jsonify({"error": "No document content was provided."}), 400

        if has_structured_content:
            docx_document = build_docx_export(title, document_data)
            buffer = io.BytesIO()
            docx_document.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"{title}.docx"
            )

        word_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{html.escape(title)}</title>
    <style>
        @page {{ size: A4 portrait; margin: 25.4mm; }}
        html, body {{
            margin: 0;
            padding: 0;
            background: #ffffff;
            color: #000000;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            line-height: 1;
        }}
        body {{
            mso-line-height-rule: exactly;
        }}
        #document-page, #result-content, #doc-header, #doc-footer,
        .section-card, .section-title, .section-body {{
            width: 100%;
            max-width: 100%;
            margin: 0;
            padding: 0;
            border: 0;
            box-shadow: none;
            background: transparent;
            color: #000000;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            line-height: 1;
        }}
        #doc-header {{
            margin-bottom: 10pt;
            min-height: 0;
        }}
        #doc-footer {{
            margin-top: 10pt;
            padding-top: 4pt;
        }}
        #doc-header:empty, #doc-footer:empty {{
            display: none;
        }}
        .section-card {{
            margin: 0 0 6pt 0;
            page-break-inside: auto;
        }}
        .section-title, .section-card h3, .section-inline-label, .inline-label {{
            display: inline;
            margin: 0;
            padding: 0;
            border: 0;
            color: #000000;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            font-weight: bold;
            line-height: 1;
            page-break-after: avoid;
        }}
        .section-body, .section-body p, #result-content p {{
            margin: 0 0 3pt 0;
            color: #000000;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            line-height: 1;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin: 0 0 3pt 0;
            padding: 0;
            border: 0;
            color: #000000;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            font-weight: bold;
            line-height: 1;
        }}
        ul, ol {{
            margin: 2pt 0 3pt 18pt;
            padding: 0;
        }}
        li {{
            margin: 0 0 1pt 0;
            line-height: 1;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            border-spacing: 0;
            margin: 0 0 4pt 0;
        }}
        th, td {{
            border: 1px solid #64748b;
            padding: 4pt 5pt;
            vertical-align: top;
            text-align: left;
            font-weight: normal;
        }}
        th {{
            font-weight: bold;
        }}
        img, svg {{
            max-width: 100%;
            height: auto;
        }}
        pre, code {{
            white-space: pre-wrap;
            word-wrap: break-word;
            font-family: "Times New Roman", Times, serif;
            font-size: 12pt;
            line-height: 1;
            background: transparent;
            border: 0;
            padding: 0;
        }}
        .print\\:hidden, .magic-btn, .magic-btn-visual, .magic-run-btn,
        button, figcaption, .absolute, .opacity-0 {{
            display: none !important;
        }}
        [contenteditable="false"] {{
            display: none;
        }}
    </style>
</head>
<body><div id="document-page">{document_html}</div></body>
</html>"""
        buffer = io.BytesIO(word_html.encode("utf-8"))
        return send_file(
            buffer,
            mimetype="application/msword",
            as_attachment=True,
            download_name=f"{title}.doc"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/section', methods=['POST'])
@require_auth
def generate_section():
    try:
        if not client:
            return jsonify({"error": "Groq API key is missing from the .env file. Please add GROQ_API_KEY."}), 500
            
        user_id = get_auth_user_id()
        ensure_credit_account(user_id)

        data = request.json
        topic = (data.get('topic', '') or '').strip()
        section_title = (data.get('section', '') or '').strip()
        is_visual = data.get('is_visual', False)
        subject_type = (data.get('type', 'Non-Coding') or 'Non-Coding').strip()
        if not section_title:
            return jsonify({"error": "Section title is required"}), 400

        system_instructions = (
            f"You are Labmate.ai, an expert academic writer. "
            f"Generate exactly one lab-practical section, not a full practical. "
            f"Topic: '{topic}'. Subject type: '{subject_type}'. Target section: '{section_title}'. "
            f"Hard boundary: every sentence must belong only to the target section. "
            f"Do not borrow content from Aim, Theory, Procedure, Code, Output, Result, or Conclusion unless that is the target section. "
            f"Do not repeat the section title as a heading or prefix (e.g., NEVER start with '{section_title}:'). Start directly with the core text. "
            f"Output clean markdown inside a strictly valid JSON object with one key, 'content'. "
            f"IMPORTANT: The 'content' value must be a standard JSON string enclosed in double quotes (\"). Do not use Python-style triple quotes (\"\"\") inside the JSON. Escape all newlines properly (\\n)."
        )
        system_instructions += " " + get_section_specific_guidance(section_title, subject_type, is_visual)
        if is_visual:
            if subject_type == 'Coding':
                system_instructions += (
                    " Output ONLY a fenced ```python``` code block inside the JSON string. "
                    "The code must generate the requested visual/output and save it with plt.savefig('plot.png'). "
                    "Do not include prose, headings, or explanations."
                )
            else:
                system_instructions += (
                    " Output ONLY one directly renderable diagram for a non-coding practical. "
                    "For UML, software design, usability design, DFD, flowchart, ER, use-case, class, sequence, state, component, deployment, and activity diagrams, return exactly one fenced ```mermaid``` block. "
                    "Use valid Mermaid syntax and short, correctly spelled node labels. "
                    "If Mermaid cannot represent the requested diagram, output clean inline SVG only. "
                    "Never output explanations, Python, imports, programming code, markdown headings, or a 'Code' section for non-coding visuals."
                )
        elif subject_type == 'Non-Coding':
            system_instructions += " This is a non-coding practical, so generally avoid source code unless the target section explicitly requires it (e.g., a 'Code' or 'Implementation' section)."

        raw_markdown_mode = is_visual or (subject_type == 'Coding' and is_code_section(section_title))
        if raw_markdown_mode:
            raw_system_instructions = (
                system_instructions
                .replace("Output clean markdown inside a strictly valid JSON object with one key, 'content'. ", "")
                .replace("IMPORTANT: The 'content' value must be a standard JSON string enclosed in double quotes (\"). Do not use Python-style triple quotes (\"\"\") inside the JSON. Escape all newlines properly (\\n).", "")
                .replace("Ensure the generated JSON is strictly valid. Do not use triple quotes inside the JSON object.", "")
                .replace(" inside the JSON string", "")
            )
            response = client.chat.completions.create(
                model=MODEL_ID,
                messages=[
                    {"role": "system", "content": raw_system_instructions},
                    {"role": "user", "content": (
                        f"Write only the '{section_title}' section for this practical topic: '{topic}'. "
                        "Return only the final markdown content. Do not wrap it in JSON."
                    )}
                ],
                temperature=0.35
            )
            content = {"content": response.choices[0].message.content}
        else:
            response = client.chat.completions.create(
                model=MODEL_ID,
                messages=[
                    {"role": "system", "content": system_instructions},
                    {"role": "user", "content": (
                        f"Write only the '{section_title}' section for this practical topic: '{topic}'. "
                        f"Return JSON only, with the shape {{\"content\": \"...\"}}."
                    )}
                ],
                response_format={"type": "json_object"},
                temperature=0.4
            )
            content = parse_model_json(response.choices[0].message.content, default_key="content")
            if "content" not in content:
                content = {"content": content.get("text", "")}

        content["content"] = clean_generated_section_content(
            content.get("content", ""),
            section_title,
            preserve_code=raw_markdown_mode
        )
        
        return jsonify(content)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/api/run_code', methods=['POST'])
@require_auth
def run_code():
    try:
        data = request.json
        code_string = data.get('code', '')

        # Security Check
        try:
            parsed = ast.parse(code_string)
            for node in ast.walk(parsed):
                if isinstance(node, (ast.Import, ast.ImportFrom)):
                    # Allow matplotlib, numpy, pandas, sklearn. Block the rest for safety.
                    allowed = ['matplotlib', 'numpy', 'pandas', 'sklearn', 'seaborn', 'plt', 'math']
                    module = node.names[0].name if isinstance(node, ast.Import) else node.module
                    if module.split('.')[0] not in allowed:
                         return jsonify({"error": f"Security Violation: '{module}' is not allowed."}), 403
        except SyntaxError as e:
            return jsonify({"error": f"Syntax Error: {str(e)}"}), 400

        sandbox_id = str(uuid.uuid4())
        sandbox_dir = os.path.join(OUTPUTS_DIR, sandbox_id)
        os.makedirs(sandbox_dir, exist_ok=True)
        
        script_path = os.path.join(sandbox_dir, 'main.py')
        # FIX 3: Ensure matplotlib doesn't try to open a GUI window (which crashes servers)
        clean_code = "import matplotlib\nmatplotlib.use('Agg')\n" + code_string
        
        with open(script_path, 'w', encoding='utf-8') as f:
            f.write(clean_code)
            
        # FIX 4: Use 'python' for Windows compatibility in venv
        result = subprocess.run(
            ['python', 'main.py'],
            cwd=sandbox_dir,
            capture_output=True,
            text=True,
            timeout=15 
        )
        
        generated_images = []
        for file in os.listdir(sandbox_dir):
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                # Adjusted URL to match the helper route
                generated_images.append(f"/outputs/{sandbox_id}/{file}")
                
        return jsonify({
            "stdout": result.stdout,
            "stderr": result.stderr,
            "images": generated_images
        })

    except subprocess.TimeoutExpired:
        return jsonify({"error": "Timeout"}), 408
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000, use_reloader=False)
